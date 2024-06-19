# Copyright (C) 2018-2022 Battelle Memorial Institute
# file: PNGDocx.py
""" Gathers all PNG files into a Word document.

After the script runs, open the DocX file in Word, select
all, then Ctrl-F9 to update fields. The figure numbers, initially
blank, should then appear.

Public Functions:
  :main: does the work
"""

import sys
import glob
import csv
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def MarkIndexEntry(entry,paragraph):
  run = paragraph.add_run()
  r = run._r
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r.append(fldChar)

  run = paragraph.add_run()
  r = run._r
  instrText = OxmlElement('w:instrText')
  instrText.set(qn('xml:space'), 'preserve')
  instrText.text = ' XE "%s" '%(entry)
  r.append(instrText)

  run = paragraph.add_run()
  r = run._r
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r.append(fldChar)

def Figure(paragraph):
  run = paragraph.add_run()
  r = run._r
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r.append(fldChar)
  instrText = OxmlElement('w:instrText')
  instrText.text = ' SEQ Figure \* ARABIC'
  r.append(instrText)
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r.append(fldChar)

def Table(paragraph):
  run = run = paragraph.add_run()
  r = run._r
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'begin')
  r.append(fldChar)
  instrText = OxmlElement('w:instrText')
  instrText.text = ' SEQ Table \* ARABIC'
  r.append(instrText)
  fldChar = OxmlElement('w:fldChar')
  fldChar.set(qn('w:fldCharType'), 'end')
  r.append(fldChar)

# python PNGDocx.py fname.docx [wide]

bWide = False
fdocx = 'Plots.docx'
if len(sys.argv) > 1:
  fdocx = sys.argv[1]
if len(sys.argv) > 2:
  bWide = bool(int(sys.argv[2]))

document = Document()
for section in document.sections:
  section.orientation = WD_ORIENT.LANDSCAPE  # this doesn't seem to do anything
  if bWide:
    section.page_width = Inches(11.0)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
  else:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.left_margin = Inches(1.00)
    section.right_margin = Inches(1.00)
  section.top_margin = Inches(0.75)
  section.bottom_margin = Inches(0.75)

#document.add_paragraph('Table of Figures')

files = sorted(glob.glob ('*.png'))
fignum = 1
for fname in files:
  if (fignum % 2) > 0:
    pass
#    document.add_page_break()
  if bWide:
    document.add_picture(fname, width=Inches(9.5))
  else:
    document.add_picture(fname, width=Inches(6.5))
# adding a cross-reference enabled caption
#  paragraph = document.add_paragraph('Figure ', style='Caption')
#  Figure (paragraph)
#  paragraph.add_run(': {:s}'.format (fname))
  fignum += 1
#  if fignum > 10:
#    break

document.save(fdocx)
